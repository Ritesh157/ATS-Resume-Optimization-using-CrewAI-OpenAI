'''
Creating utils.py file to:
    1) store reusable functions that are used in many places in project.
    2) Helper Code (small task): reading files, converting formats, cleaning text, saving files

'''
from docx import Document # This lets you create and edit .docx Word files
from io import BytesIO # Creates a “fake file” in memory (without saving to disk).

'''
This function takes plain text and converts it into a .docx file, and returns it as bytes.
This is helpful when you want to:
    1) Return a file in a web app (like Streamlit)
    2) Send the file to a user
'''
def txt_to_docx_bytes(text: str) -> bytes:
    doc = Document() # This creates a new blank Word document
    for line in text.splitlines(): # Splits the text into lines (wherever there is a newline \n). Goes through each line one by one.
        if line.strip == "": # If the line is empty → add an empty paragraph (blank line)
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line) # If the line has text → add it as a paragraph in the DOCX
    out = BytesIO() # Creates a “fake file” in memory
    doc.save(out) # # Saves the DOCX content into it
    return out.getvalue() # Returns the file content in bytes format. These bytes can then be: downloaded, emailed, sent to another agent, saves later as a .docx file